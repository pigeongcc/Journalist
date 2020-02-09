# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'c:\Users\gulya\Desktop\Journalooger\design-1.ui'
#
# Created by: PyQt5 UI code generator 5.6
#
# WARNING! All changes made in this file will be lost!

import sys
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QFileDialog
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
from docx.shared import Pt
#import logic

class Ui_Dialog(object):
    def __init__(self):
        super().__init__()

        self.mainwindow = QtWidgets.QMainWindow()
        self.setupUi(self.mainwindow)
        self.mainwindow.show()

    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.resize(277, 182)
        self.label = QtWidgets.QLabel(Dialog)
        self.label.setGeometry(QtCore.QRect(20, 10, 251, 61))
        font = QtGui.QFont()
        font.setFamily("MV Boli")
        font.setPointSize(22)
        font.setBold(True)
        font.setWeight(75)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.quitButton = QtWidgets.QPushButton(Dialog)
        self.quitButton.setGeometry(QtCore.QRect(90, 120, 93, 28))
        self.quitButton.setObjectName("quitButton")
        self.quitButton.clicked.connect(self.mainwindow.close)
        self.startButton = QtWidgets.QPushButton(Dialog)
        self.startButton.setGeometry(QtCore.QRect(90, 80, 93, 28))
        self.startButton.setObjectName("startButton")
        self.startButton.clicked.connect(self.open_file)

        self.retranslateUi(Dialog)
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def open_file(self):
        filename = QFileDialog.getOpenFileName(self.mainwindow, "Открыть файл Word", '/home')
        if filename[0]:
            self.FILENAME = filename[0]
            self.word = Document(self.FILENAME)
            self.calculations()


    def calculations(self):
        wordTable = self.word.tables[0]
        self.quantityOfColumns = len(wordTable.rows[0].cells)
        self.quantityOfDisciplines = self.quantityOfColumns - 4
        self.quantityOfLines = len(wordTable.rows)
        self.quantityOfCadets = self.quantityOfLines - 2

        #AVGes of Cadets
        summGlobal = 0
        quantityGlobal = 0
        for i in range(1, 1+self.quantityOfCadets):
            summ = ""
            for j in range(1+2, 1+self.quantityOfDisciplines+2):
                cellText = wordTable.rows[i].cells[j].text
                summ += str(cellText)

            try:
                summ, quantity = self.stringGradesToSum(summ) # sum == "354" -> sum == 12
                avg = round(summ / quantity, 2)
                wordTable.rows[i].cells[1+self.quantityOfDisciplines+2].text = str(avg)
                summGlobal += summ
                quantityGlobal += quantity
            except Exception as e:
                print(e, " after AVGs of Cadets")

        try:
            #Global AVG
            avgGlobal = round(summGlobal / quantityGlobal, 2)
            wordTable.rows[1+self.quantityOfCadets].cells[1+self.quantityOfDisciplines+2].text = str(avgGlobal)
        except Exception as e:
            print(e, " in Global AVG")

        #AVGes of Disciplines
        for j in range(1+2, 1+self.quantityOfDisciplines+2):
            summ = ""
            for i in range(1, 1+self.quantityOfCadets):
                cellText = wordTable.rows[i].cells[j].text
                summ += str(cellText)

            try:
                summ, quantity = self.stringGradesToSum(summ) # sum == "354" -> sum == 12
                avg = round(summ / quantity, 2)
                wordTable.rows[1+self.quantityOfCadets].cells[j].text = str(avg)
            except Exception as e:
                print(e, " after AVGs of Disc-s")

        self.save_word()
    
    """ August 03, 2019
    I've stopped at the moment when I got two-dimension list GRADES with all the grades from the table
    But the list has additional, empty cells. Next step is to calculate averages.
    """

    def save_word(self):
        option = QFileDialog.Options()
        fname, _ = QFileDialog.getSaveFileName(None, "QFileDialog.getSaveFileName()", "", "Microsoft Word Documents (*.docx);;All Files (*)")
        if fname:
            self.word.save(fname)
            self.word = Document(fname)
            self.word.save(fname)
            print("before Styling")
            self.style_changing()
            print("after Styling")
            self.word.save(fname)

    def stringGradesToSum(self, sumString):
        length = len(sumString)
        sumInt = 0
        for i in range(length):
            sumInt += int(sumString[0])
            sumString = sumString[1::]
        return sumInt, length

    def style_changing(self):

        # HEIGHT
        wordTable = self.word.tables[0]
        for i in range(1, self.quantityOfColumns - 2):
            wordTable.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            wordTable.rows[i].height = Cm(0.47)
            wordTable.allow_autofit = True

        # FONT and SIZE

                # surnames & numbers
        for row in wordTable.rows[0:1+self.quantityOfCadets]:
            """
            for cell in column.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.get_or_add_tcW()
                tcW.type = 'auto'
            """
            for cell in row.cells[0:3]:
                try:                    
                    paragraph = cell.paragraphs[0]
                    run = paragraph.runs
                    font = run[0].font
                    font.name = 'Times New Roman'
                    font.size = Pt(10)
                except Exception as e:
                    print('S1 ', e)

                # headings (with names of disciplines)
        row = wordTable.rows[0]
        for cell in row.cells:
            try:
                paragraph = cell.paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run = paragraph.runs
                font = run[0].font
                font.name = 'Times New Roman'
                font.size = Pt(9)
            except Exception as e:
                print('S2 ', e)

                # grades
        for row in wordTable.rows[1:1+self.quantityOfCadets]:
            for cell in row.cells[1+2:1+self.quantityOfDisciplines+2]:
                try:
                    paragraph = cell.paragraphs[0]
                    run = paragraph.runs
                    font = run[0].font
                    font.name = 'Calibri'
                    font.size = Pt(9)
                except Exception as e:
                    print('S3 ', e)

                # avg Disciplines
        row = wordTable.rows[1+self.quantityOfCadets]
        for cell in row.cells[1+2:1+self.quantityOfDisciplines+2 + 1]:
            try:
                paragraph = cell.paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run = paragraph.runs
                font = run[0].font
                font.name = 'Times New Roman'
                font.size = Pt(10)
                font.bold = True
            except Exception as e:
                print('S4 ', e)

                # avg Cadets
        for row in wordTable.rows[1:1+self.quantityOfCadets]:
            cell = row.cells[1+self.quantityOfDisciplines+2]
            try:
                paragraph = cell.paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run = paragraph.runs
                font = run[0].font
                font.name = 'Times New Roman'
                font.size = Pt(10)
                font.bold = True
            except Exception as e:
                print('S5 ', e)

        #              rows[26] == rows[1+self.quantityOfCadets]             cells[self.N] == cells[1+self.quantityOfDisciplines+2]
        paragraph = wordTable.rows[1+self.quantityOfCadets].cells[1+self.quantityOfDisciplines+2].paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        wordTable.cell(1+self.quantityOfCadets, 1+self.quantityOfDisciplines+2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph.space_after = Pt(0)
        run = paragraph.runs
        try:
            font = run[0].font
            font.size = Pt(12)
        except Exception as e:
            print('S6 ', e)

        

        for row in wordTable.rows[0:1+self.quantityOfCadets]:
            for cell in row.cells[0:1+self.quantityOfDisciplines+2]:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.get_or_add_tcW()
                tcW.type = 'auto'

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Dialog"))
        self.label.setText(_translate("Dialog", "Journalooger"))
        self.quitButton.setText(_translate("Dialog", "Конец"))
        self.startButton.setText(_translate("Dialog", "Старт"))

